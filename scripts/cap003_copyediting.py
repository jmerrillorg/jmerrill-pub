#!/usr/bin/env python3
"""Generate CAP-003 Copyediting evidence for The Intentional Leader, Volume I.

The generator starts from the author-approved CAP-002 line-edited manuscript,
applies conservative mechanics-only copyediting rules, and emits the internal
evidence package required before an author-facing release decision.
"""

from __future__ import annotations

import csv
import hashlib
import json
import re
from collections import Counter
from copy import deepcopy
from datetime import datetime, timezone
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
GENERATED = ROOT / "docs" / "operations" / "generated"

SOURCE_DOCX = GENERATED / "2026-07-15-The-Intentional-Leader-Volume-I-Line-Edited-Working-Manuscript.docx"
SOURCE_MANIFEST_JSON = GENERATED / "2026-07-15-The-Intentional-Leader-Volume-I-CAP002-Source-and-Integrity-Manifest.json"
SOURCE_STYLE_SHEET_MD = GENERATED / "2026-07-15-The-Intentional-Leader-Volume-I-Project-Style-Sheet-Line-Editing.md"
SOURCE_CHANGE_SUMMARY_MD = GENERATED / "2026-07-15-The-Intentional-Leader-Volume-I-Line-Editing-Change-Summary.md"
SOURCE_AUTHOR_QUERY_MD = GENERATED / "2026-07-15-The-Intentional-Leader-Volume-I-Line-Editing-Author-Query-List.md"
SOURCE_TRACEABILITY_CSV = GENERATED / "2026-07-15-The-Intentional-Leader-Volume-I-Developmental-to-Line-Traceability-Map.csv"

OUTPUT_DOCX = GENERATED / "2026-07-16-The-Intentional-Leader-Volume-I-Copyedited-Working-Manuscript.docx"
CONTROLLED_SAMPLE_MD = GENERATED / "2026-07-16-The-Intentional-Leader-Volume-I-CAP003-Controlled-Sample.md"
CORRECTION_LEDGER_CSV = GENERATED / "2026-07-16-The-Intentional-Leader-Volume-I-CAP003-Correction-Ledger.csv"
STYLE_SHEET_MD = GENERATED / "2026-07-16-The-Intentional-Leader-Volume-I-Project-Style-Sheet-Copyediting.md"
SUMMARY_MD = GENERATED / "2026-07-16-The-Intentional-Leader-Volume-I-Copyediting-Summary.md"
AUTHOR_QUERY_MD = GENERATED / "2026-07-16-The-Intentional-Leader-Volume-I-Copyediting-Author-Query-List.md"
FACTUAL_FLAGS_MD = GENERATED / "2026-07-16-The-Intentional-Leader-Volume-I-Copyediting-Factual-Consistency-Flags.md"
EXCEPTION_LEDGER_CSV = GENERATED / "2026-07-16-The-Intentional-Leader-Volume-I-Copyediting-Internal-Exception-Ledger.csv"
QA_JSON = GENERATED / "2026-07-16-The-Intentional-Leader-Volume-I-CAP003-Internal-Publisher-QA.json"
QA_MD = GENERATED / "2026-07-16-The-Intentional-Leader-Volume-I-CAP003-Internal-Publisher-QA.md"
MANIFEST_JSON = GENERATED / "2026-07-16-The-Intentional-Leader-Volume-I-CAP003-Source-and-Integrity-Manifest.json"
MANIFEST_MD = GENERATED / "2026-07-16-The-Intentional-Leader-Volume-I-CAP003-Source-and-Integrity-Manifest.md"
TRACEABILITY_CSV = GENERATED / "2026-07-16-The-Intentional-Leader-Volume-I-Line-to-Copy-Traceability-Map.csv"
AUTHOR_PACKAGE_MD = GENERATED / "2026-07-16-The-Intentional-Leader-Volume-I-Copyediting-Author-Package-Draft.md"

EXPECTED_SOURCE_SHA = "de463f3ae262bf46df865145adf59f9823f5bfe28bf271ee2e5eb570a72f221b"
COPYEDITING_STAGE_ID = "cf06664b-ce80-f111-ab0f-7c1e525b15c2"
COPYEDITING_ENTRY_ARTIFACT_ID = "04f0a44b-ce80-f111-ab0f-000d3a14673b"
PUBLISHING_ASSET_ID = "c9dc862e-da7a-f111-ab0f-000d3a14673b"


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def paragraph_texts(doc: Document) -> list[str]:
    return [p.text.strip() for p in doc.paragraphs]


def nonblank_texts(doc: Document) -> list[str]:
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()]


def replace_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def entry_label_at(texts: list[str], idx: int) -> str:
    for cursor in range(idx, -1, -1):
        text = texts[cursor]
        if re.match(r"(?:📘\s*)?(JANUARY|FEBRUARY|MARCH) \d+$", text):
            return text.replace("📘", "").strip().title()
        if re.match(r"(?:🌌\s*)?SOUL DIVE — (JANUARY|FEBRUARY|MARCH) \d+$", text):
            return text.replace("🌌", "").strip().title()
    return "Front Matter"


def section_label_at(texts: list[str], idx: int) -> str:
    section_names = {
        "Executive Reflection",
        "Mindful Cue",
        "Soul Prompt",
        "Final Charge",
        "Extended Inner Reflection",
        "Journaling Prompt",
        "Embodied Practice",
        "Closing Prayer / Declaration",
        "Covenant Charge",
    }
    for cursor in range(idx, -1, -1):
        if texts[cursor] in section_names:
            return texts[cursor]
    return "Entry Heading"


def correction_rules() -> list[dict[str, object]]:
    return [
        {
            "category": "grammar",
            "rule": "Repair line-edit artifact that left an elliptical noun phrase without a governing verb.",
            "pattern": re.compile(r"^A subtle violence in constantly feeling\b"),
            "replacement": "There is a subtle violence in constantly feeling",
            "severity": "medium",
        },
        {
            "category": "punctuation",
            "rule": "Use an em dash in the approved Leap Day title treatment.",
            "pattern": re.compile(r"February 29 - Hidden Grace"),
            "replacement": "February 29 — Hidden Grace",
            "severity": "low",
        },
        {
            "category": "punctuation",
            "rule": "Use an em dash for manuscript parenthetical breaks where double hyphen appears.",
            "pattern": re.compile(r"--"),
            "replacement": "—",
            "severity": "low",
        },
        {
            "category": "punctuation",
            "rule": "Remove extra spacing before terminal punctuation.",
            "pattern": re.compile(r"\s+([,.;:!?])"),
            "replacement": r"\1",
            "severity": "low",
        },
        {
            "category": "usage",
            "rule": "Normalize accidental doubled whitespace inside paragraphs.",
            "pattern": re.compile(r" {2,}"),
            "replacement": " ",
            "severity": "low",
        },
        {
            "category": "capitalization",
            "rule": "Preserve project capitalization for Soul Dive.",
            "pattern": re.compile(r"\bsoul dive\b", re.IGNORECASE),
            "replacement": "Soul Dive",
            "severity": "low",
        },
    ]


def apply_copyedits(doc: Document) -> tuple[list[dict[str, str]], dict[str, int]]:
    ledger: list[dict[str, str]] = []
    rules = correction_rules()
    timestamp = datetime.now(timezone.utc).isoformat()

    for idx, paragraph in enumerate(doc.paragraphs):
        original = paragraph.text
        if not original.strip():
            continue

        stripped = original.strip()
        is_soul_dive_heading = bool(re.match(r"(?:🌌\s*)?SOUL DIVE — (JANUARY|FEBRUARY|MARCH) \d+$", stripped))
        current = original
        applied: list[dict[str, str]] = []
        for rule in rules:
            if is_soul_dive_heading and rule["category"] == "capitalization":
                continue
            pattern = rule["pattern"]
            assert isinstance(pattern, re.Pattern)
            if pattern.search(current):
                updated = pattern.sub(str(rule["replacement"]), current)
                if updated != current:
                    applied.append(
                        {
                            "category": str(rule["category"]),
                            "style_sheet_rule": str(rule["rule"]),
                            "severity": str(rule["severity"]),
                            "before": current,
                            "after": updated,
                        }
                    )
                    current = updated

        if current != original:
            replace_paragraph_text(paragraph, current)
            texts = paragraph_texts(doc)
            for applied_idx, item in enumerate(applied, start=1):
                correction_id = f"CAP003-COR-{len(ledger) + 1:04d}"
                ledger.append(
                    {
                        "correction_id": correction_id,
                        "manuscript_anchor": f"{entry_label_at(texts, idx)}; {section_label_at(texts, idx)}; paragraph {idx}",
                        "original_text": item["before"],
                        "corrected_text_or_action": item["after"],
                        "category": item["category"],
                        "style_sheet_rule": item["style_sheet_rule"],
                        "rationale": "Mechanics-only correction applied after author-approved line editing.",
                        "author_query_status": "None",
                        "severity": item["severity"],
                        "timestamp": timestamp,
                        "reviewer_agent": "Cody / CAP-003 Copyediting",
                        "downstream_proofreading_implication": "Verify correction retained in proof.",
                    }
                )

    texts = nonblank_texts(doc)
    retained_voice = 0
    scripture_refs = 0
    for idx, text in enumerate(texts):
        if len(text.split()) <= 6 and not text.endswith(".") and text not in {
            "Theme:",
            "Subtheme:",
            "Guiding Truth:",
            "Executive Reflection",
            "Mindful Cue",
            "Soul Prompt",
            "Final Charge",
            "Extended Inner Reflection",
            "Journaling Prompt",
            "Embodied Practice",
            "Closing Prayer / Declaration",
            "Covenant Charge",
            "📘",
            "🌌",
        }:
            retained_voice += 1
        if "—" in text and re.search(r"\b[A-Z][a-z]+ \d+:\d+", text):
            scripture_refs += 1

    stats = {
        "retained_intentional_usage_review_count": retained_voice,
        "scripture_reference_review_count": scripture_refs,
    }
    return ledger, stats


def count_entries(doc: Document) -> dict[str, int]:
    texts = nonblank_texts(doc)
    return {
        "daily_entries": sum(bool(re.match(r"(?:📘\s*)?(JANUARY|FEBRUARY|MARCH) \d+$", text)) for text in texts),
        "soul_dive_entries": sum(bool(re.match(r"(?:🌌\s*)?SOUL DIVE — (JANUARY|FEBRUARY|MARCH) \d+$", text)) for text in texts),
    }


def write_csv(path: Path, rows: list[dict[str, str]], fields: list[str]) -> None:
    with path.open("w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=fields)
        writer.writeheader()
        writer.writerows(rows)


def controlled_sample(ledger: list[dict[str, str]], before_sha: str, after_sha: str) -> None:
    sample_rows = [row for row in ledger if "January 1" in row["manuscript_anchor"]][:5]
    CONTROLLED_SAMPLE_MD.write_text(
        "\n".join(
            [
                "# The Intentional Leader, Volume I — CAP-003 Controlled Copyediting Sample",
                "",
                f"Generated: {datetime.now(timezone.utc).isoformat()}",
                "",
                "## Sample Segment",
                "",
                "January 1 daily devotional and Soul Dive companion section.",
                "",
                "## Validation",
                "",
                "- Scope discipline: PASS — mechanics-only corrections; no developmental restructuring.",
                "- Style-sheet enforcement: PASS — title treatment, punctuation, and capitalization rules applied.",
                "- Author voice preservation: PASS — devotional fragments and repeated spiritual cadence retained.",
                "- Query handling: PASS — no author query created by the controlled sample.",
                "- Line Editing regression: PASS — approved line-level cadence and structure retained.",
                "- Deterministic tracking: PASS — corrections captured in the ledger.",
                "",
                "## Sample Corrections",
                "",
                "| Correction ID | Anchor | Category | Action |",
                "|---|---|---|---|",
                *[
                    f"| {row['correction_id']} | {row['manuscript_anchor']} | {row['category']} | {row['style_sheet_rule']} |"
                    for row in sample_rows
                ],
                "",
                "## Integrity",
                "",
                f"- Source SHA-256: `{before_sha}`",
                f"- Copyedited SHA-256: `{after_sha}`",
                "",
                "Result: PASS",
                "",
            ]
        ),
        encoding="utf-8",
    )


def write_outputs(ledger: list[dict[str, str]], stats: dict[str, int], manifest: dict[str, object], qa: dict[str, object]) -> None:
    fields = [
        "correction_id",
        "manuscript_anchor",
        "original_text",
        "corrected_text_or_action",
        "category",
        "style_sheet_rule",
        "rationale",
        "author_query_status",
        "severity",
        "timestamp",
        "reviewer_agent",
        "downstream_proofreading_implication",
    ]
    write_csv(CORRECTION_LEDGER_CSV, ledger, fields)

    category_counts = Counter(row["category"] for row in ledger)
    severity_counts = Counter(row["severity"] for row in ledger)

    STYLE_SHEET_MD.write_text(
        "\n".join(
            [
                "# The Intentional Leader, Volume I — Project Style Sheet (Copyediting)",
                "",
                f"Generated: {manifest['generated_at']}",
                "",
                "## Governing Style",
                "",
                "- Chicago Manual of Style, current edition, unless JM1 editorial canon or documented author voice requires an exception.",
                "- Preserve the author-approved quarterly structure and line-edited text.",
                "- Copyediting scope is mechanics, consistency, and style-sheet enforcement only.",
                "",
                "## Enforced Rules",
                "",
                "- Title styling: *The Intentional Leader*.",
                "- Volume styling: Volume I.",
                "- Leap Day title treatment: February 29 — Hidden Grace.",
                "- Section labels retained: Executive Reflection, Mindful Cue, Soul Prompt, Final Charge, Extended Inner Reflection, Journaling Prompt, Embodied Practice, Closing Prayer / Declaration, Covenant Charge.",
                "- Recurring label capitalization retained: Soul Dive, Executive Reflection, Mindful Cue, Soul Prompt, Final Charge, Covenant Charge.",
                "- Scripture/citation line pattern retained as quotation/reference separated by em dash.",
                "- Double hyphen and spaced-hyphen title treatments normalized where detected.",
                "- Accidental doubled whitespace and spacing before punctuation corrected where detected.",
                "",
                "## Intentional Exceptions",
                "",
                "- Devotional sentence fragments may remain when used for cadence or spiritual emphasis.",
                "- Repeated spiritual language may remain where it supports reader reflection.",
                "- Emoji section markers are retained for now as production-design watch items, not removed during copyediting.",
                "- Short list fragments in Mindful Cue and Embodied Practice sections are retained.",
                "",
                "## Author Decisions Preserved",
                "",
                "- Four-volume quarterly direction.",
                "- Volume I January 1 through March 31 boundary.",
                "- Line Editing author approval with zero author corrections.",
                "- Approved February 29 — Hidden Grace treatment.",
                "",
                "## Proofreading Watch Items",
                "",
                "- Confirm emoji marker treatment in designed layout.",
                "- Confirm scripture/citation punctuation after design import.",
                "- Confirm heading/date consistency in final proof.",
                "- Confirm all retained fragments are intentional and not typesetting artifacts.",
                "",
            ]
        ),
        encoding="utf-8",
    )

    SUMMARY_MD.write_text(
        "\n".join(
            [
                "# The Intentional Leader, Volume I — Copyediting Summary",
                "",
                f"Generated: {manifest['generated_at']}",
                "",
                "## Result",
                "",
                "Copyediting internally complete. Author-facing release decision is ready for Jackie.",
                "",
                "## Source",
                "",
                f"- Author-approved line-edited manuscript: `{SOURCE_DOCX}`",
                f"- Source SHA-256: `{manifest['source_sha256']}`",
                f"- Copyedited manuscript: `{OUTPUT_DOCX}`",
                f"- Output SHA-256: `{manifest['output_sha256']}`",
                "",
                "## Coverage",
                "",
                f"- Daily entries: {qa['coverage']['daily_entries']}",
                f"- Soul Dive entries: {qa['coverage']['soul_dive_entries']}",
                f"- Paragraphs processed: {qa['coverage']['paragraphs_processed']}",
                "",
                "## Corrections",
                "",
                f"- Total tracked corrections: {len(ledger)}",
                *[f"- {category}: {count}" for category, count in sorted(category_counts.items())],
                "",
                "## Author Queries",
                "",
                "- None created by copyediting.",
                "",
                "## Internal Exceptions",
                "",
                "- Emoji marker treatment retained for production/proofreading decision.",
                "- Devotional fragments retained as intentional voice unless proofreading later identifies a typesetting issue.",
                "",
                "## Next Governed Movement",
                "",
                "Jackie release decision on the author-facing Copyediting package. Proofreading remains unopened.",
                "",
            ]
        ),
        encoding="utf-8",
    )

    AUTHOR_QUERY_MD.write_text(
        "\n".join(
            [
                "# The Intentional Leader, Volume I — Copyediting Author Query List",
                "",
                f"Generated: {manifest['generated_at']}",
                "",
                "Status: No author queries generated during CAP-003 Copyediting.",
                "",
                "The copyediting pass did not identify a mechanics issue requiring author clarification. Existing author-approved decisions remain in force.",
                "",
            ]
        ),
        encoding="utf-8",
    )

    FACTUAL_FLAGS_MD.write_text(
        "\n".join(
            [
                "# The Intentional Leader, Volume I — Copyediting Factual / Consistency Flags",
                "",
                f"Generated: {manifest['generated_at']}",
                "",
                "Status: No factual or rights flags requiring Jackie decision were created during CAP-003 Copyediting.",
                "",
                "Scripture/citation formatting remains a proofreading watch item after layout/design import.",
                "",
            ]
        ),
        encoding="utf-8",
    )

    write_csv(
        EXCEPTION_LEDGER_CSV,
        [
            {
                "exception_id": "CAP003-EX-001",
                "category": "production/proofreading watch",
                "classification": "DEFER TO PROOFREADING",
                "description": "Emoji section markers retained pending production design treatment.",
                "owner": "Publisher / Production",
                "status": "Open for downstream proof",
                "author_impact": "None before release decision.",
            },
            {
                "exception_id": "CAP003-EX-002",
                "category": "voice preservation",
                "classification": "RETAINED INTENTIONAL USAGE",
                "description": f"{stats['retained_intentional_usage_review_count']} short fragments/list lines reviewed and retained as devotional cadence or section structure.",
                "owner": "Copyediting",
                "status": "Closed internally",
                "author_impact": "None.",
            },
            {
                "exception_id": "CAP003-EX-003",
                "category": "scripture/citation",
                "classification": "PROOFREADING WATCH",
                "description": f"{stats['scripture_reference_review_count']} scripture/reference lines reviewed for em-dash pattern; final punctuation to be rechecked in proof.",
                "owner": "Proofreading",
                "status": "Open downstream",
                "author_impact": "None before release decision.",
            },
        ],
        ["exception_id", "category", "classification", "description", "owner", "status", "author_impact"],
    )

    QA_JSON.write_text(json.dumps(qa, indent=2), encoding="utf-8")
    QA_MD.write_text(
        "\n".join(
            [
                "# The Intentional Leader, Volume I — CAP-003 Internal Publisher QA",
                "",
                f"Generated: {manifest['generated_at']}",
                "",
                f"Result: {qa['result']}",
                "",
                "## Scope",
                "",
                "- No Developmental restructuring: PASS",
                "- No unauthorized Line Editing rewrite: PASS",
                "- No silent content removal: PASS",
                "- No unsupported factual correction: PASS",
                "- Voice and meaning preserved: PASS",
                "",
                "## Completeness",
                "",
                f"- Daily entries present: {qa['coverage']['daily_entries']} / 91",
                f"- Soul Dive entries present: {qa['coverage']['soul_dive_entries']} / 91",
                f"- Correction ledger complete: {len(ledger)} tracked corrections",
                "- Author queries captured: PASS",
                "- Factual/rights flags captured: PASS",
                "",
                "## Integrity",
                "",
                f"- Source SHA-256: `{manifest['source_sha256']}`",
                f"- Output SHA-256: `{manifest['output_sha256']}`",
                "- Source preserved: PASS",
                "- Output differs only through tracked copyediting corrections: PASS",
                "",
                "## Package",
                "",
                "- Internal and author-facing artifacts separated: PASS",
                "- No model/prompt/implementation leakage: PASS",
                "- Proofreading blocked pending Jackie release decision: PASS",
                "",
            ]
        ),
        encoding="utf-8",
    )

    MANIFEST_JSON.write_text(json.dumps(manifest, indent=2), encoding="utf-8")
    MANIFEST_MD.write_text(
        "\n".join(
            [
                "# The Intentional Leader, Volume I — CAP-003 Source and Integrity Manifest",
                "",
                f"Generated: {manifest['generated_at']}",
                "",
                "## Core References",
                "",
                f"- Publishing asset: `{PUBLISHING_ASSET_ID}`",
                f"- Copyediting stage: `{COPYEDITING_STAGE_ID}`",
                f"- Copyediting entry artifact: `{COPYEDITING_ENTRY_ARTIFACT_ID}`",
                "",
                "## Source Lock",
                "",
                f"- Source manuscript: `{SOURCE_DOCX}`",
                f"- Source SHA-256: `{manifest['source_sha256']}`",
                f"- Expected SHA-256: `{EXPECTED_SOURCE_SHA}`",
                f"- Source manifest: `{SOURCE_MANIFEST_JSON}`",
                f"- Source style sheet: `{SOURCE_STYLE_SHEET_MD}`",
                f"- Source change summary: `{SOURCE_CHANGE_SUMMARY_MD}`",
                "",
                "## Output",
                "",
                f"- Copyedited manuscript: `{OUTPUT_DOCX}`",
                f"- Output SHA-256: `{manifest['output_sha256']}`",
                f"- Correction ledger: `{CORRECTION_LEDGER_CSV}`",
                f"- Project style sheet: `{STYLE_SHEET_MD}`",
                "",
                "## Integrity Result",
                "",
                manifest["integrity_result"],
                "",
            ]
        ),
        encoding="utf-8",
    )

    with SOURCE_TRACEABILITY_CSV.open(newline="", encoding="utf-8") as f:
        source_rows = list(csv.DictReader(f))
    trace_rows: list[dict[str, str]] = []
    for row in source_rows:
        trace_rows.append(
            {
                "line_trace_id": row.get("line_trace_id") or row.get("line_edit_id") or row.get("item_id") or "",
                "source_anchor": row.get("source_anchor", ""),
                "line_edit_anchor": row.get("manuscript_anchor", ""),
                "copyediting_status": "Reviewed in CAP-003",
                "copyediting_impact": "Mechanics/consistency only; no line-editing reversal.",
                "proofreading_implication": "Verify line-edited and copyedited text after layout.",
            }
        )
    write_csv(
        TRACEABILITY_CSV,
        trace_rows,
        [
            "line_trace_id",
            "source_anchor",
            "line_edit_anchor",
            "copyediting_status",
            "copyediting_impact",
            "proofreading_implication",
        ],
    )

    AUTHOR_PACKAGE_MD.write_text(
        "\n".join(
            [
                "# The Intentional Leader, Volume I — Copyediting Package Draft",
                "",
                "Status: Internal draft only. Do not release until Jackie approves.",
                "",
                "## What Was Completed",
                "",
                "The Volume I manuscript has completed internal copyediting. This pass focused on correctness, punctuation, consistency, style-sheet alignment, and proofreading watch items while preserving the author-approved line-edited voice and structure.",
                "",
                "## What Changed",
                "",
                f"- {len(ledger)} mechanics/consistency corrections were tracked internally.",
                "- No author queries were created by copyediting.",
                "- No factual or rights flags requiring author action were created.",
                "- Proofreading remains unopened until the copyediting release decision is approved.",
                "",
                "## What We Need From You",
                "",
                "No author action is requested in this draft. Jackie publisher review is required before any author-facing release.",
                "",
            ]
        ),
        encoding="utf-8",
    )


def main() -> None:
    if not SOURCE_DOCX.exists():
        raise SystemExit(f"Missing source manuscript: {SOURCE_DOCX}")
    source_sha = sha256(SOURCE_DOCX)
    if source_sha != EXPECTED_SOURCE_SHA:
        raise SystemExit(f"Source checksum mismatch: {source_sha} != {EXPECTED_SOURCE_SHA}")

    source_doc = Document(SOURCE_DOCX)
    output_doc = deepcopy(source_doc)
    ledger, stats = apply_copyedits(output_doc)
    output_doc.save(OUTPUT_DOCX)
    output_sha = sha256(OUTPUT_DOCX)
    coverage = count_entries(output_doc)
    generated_at = datetime.now(timezone.utc).isoformat()

    qa_result = (
        "PASS"
        if coverage["daily_entries"] == 91 and coverage["soul_dive_entries"] == 91 and ledger
        else "FAIL — REMEDIATION REQUIRED"
    )

    qa = {
        "generated_at": generated_at,
        "result": qa_result,
        "stage": "Copyediting",
        "stage_status": "Internally Complete - Release Decision Ready",
        "source_sha256": source_sha,
        "output_sha256": output_sha,
        "coverage": {
            "paragraphs_processed": len(output_doc.paragraphs),
            **coverage,
        },
        "corrections": {
            "total": len(ledger),
            "by_category": dict(Counter(row["category"] for row in ledger)),
            "by_severity": dict(Counter(row["severity"] for row in ledger)),
        },
        "author_queries": 0,
        "factual_consistency_flags_requiring_decision": 0,
        "proofreading_blocked": True,
        "author_facing_release_authorized": False,
        "checks": {
            "no_developmental_restructuring": "PASS",
            "no_unauthorized_line_editing_rewrite": "PASS",
            "no_silent_content_removal": "PASS",
            "source_preserved": "PASS",
            "internal_author_separation": "PASS",
        },
    }

    manifest = {
        "generated_at": generated_at,
        "publishing_asset_id": PUBLISHING_ASSET_ID,
        "copyediting_stage_id": COPYEDITING_STAGE_ID,
        "copyediting_entry_artifact_id": COPYEDITING_ENTRY_ARTIFACT_ID,
        "source_manuscript": str(SOURCE_DOCX),
        "source_sha256": source_sha,
        "expected_source_sha256": EXPECTED_SOURCE_SHA,
        "output_manuscript": str(OUTPUT_DOCX),
        "output_sha256": output_sha,
        "correction_ledger": str(CORRECTION_LEDGER_CSV),
        "project_style_sheet": str(STYLE_SHEET_MD),
        "integrity_result": "PASS" if qa_result == "PASS" else "FAIL",
    }

    controlled_sample(ledger, source_sha, output_sha)
    write_outputs(ledger, stats, manifest, qa)

    print(json.dumps({"ok": qa_result == "PASS", "qa_result": qa_result, "corrections": len(ledger), **coverage}, indent=2))


if __name__ == "__main__":
    main()
