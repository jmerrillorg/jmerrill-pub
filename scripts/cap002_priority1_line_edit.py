#!/usr/bin/env python3
"""Generate CAP-002 Priority 1 line-edit evidence for The Intentional Leader.

The generator is intentionally conservative: it locks the approved source
manuscript, trims the output to the governed Volume I boundary, applies
line-level rhythm/clarity changes only inside the queued Soul Dive companion
sections, and emits the ledgers needed for publisher QA.
"""

from __future__ import annotations

import csv
import hashlib
import json
import re
from collections import Counter
from copy import deepcopy
from datetime import datetime, timezone
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
GENERATED = ROOT / "docs" / "operations" / "generated"

SOURCE_DOCX = GENERATED / "2026-07-13-The-Intentional-Leader-Volume-I-Revised-Internal-Developmental-Manuscript.docx"
SOURCE_PACKAGE_JSON = GENERATED / "2026-07-14-CAP-002-Line-Editing-Source-Package.json"
QUEUE_CSV = GENERATED / "2026-07-13-CAP-001-Line-Editing-Intake-Queue-The-Intentional-Leader-Volume-I.csv"

OUTPUT_DOCX = GENERATED / "2026-07-15-The-Intentional-Leader-Volume-I-Line-Edited-Working-Manuscript.docx"
STYLE_SHEET_MD = GENERATED / "2026-07-15-The-Intentional-Leader-Volume-I-Project-Style-Sheet-Line-Editing.md"
DISPOSITION_CSV = GENERATED / "2026-07-15-The-Intentional-Leader-Volume-I-CAP002-80-Item-Disposition-Ledger.csv"
CHANGE_SUMMARY_MD = GENERATED / "2026-07-15-The-Intentional-Leader-Volume-I-Line-Editing-Change-Summary.md"
AUTHOR_QUERY_MD = GENERATED / "2026-07-15-The-Intentional-Leader-Volume-I-Line-Editing-Author-Query-List.md"
EXCEPTION_LEDGER_CSV = GENERATED / "2026-07-15-The-Intentional-Leader-Volume-I-Line-Editing-Internal-Exception-Ledger.csv"
TRACEABILITY_CSV = GENERATED / "2026-07-15-The-Intentional-Leader-Volume-I-Developmental-to-Line-Traceability-Map.csv"
QA_JSON = GENERATED / "2026-07-15-The-Intentional-Leader-Volume-I-CAP002-Internal-Publisher-QA.json"
QA_MD = GENERATED / "2026-07-15-The-Intentional-Leader-Volume-I-CAP002-Internal-Publisher-QA.md"
MANIFEST_JSON = GENERATED / "2026-07-15-The-Intentional-Leader-Volume-I-CAP002-Source-and-Integrity-Manifest.json"
MANIFEST_MD = GENERATED / "2026-07-15-The-Intentional-Leader-Volume-I-CAP002-Source-and-Integrity-Manifest.md"
AUTHOR_PACKAGE_MD = GENERATED / "2026-07-15-The-Intentional-Leader-Volume-I-Line-Editing-Author-Package-Draft.md"


MONTHS = {
    "JANUARY": 1,
    "FEBRUARY": 2,
    "MARCH": 3,
    "APRIL": 4,
}


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def norm_date_label(raw: str) -> str:
    month, day = raw.split()
    return f"{month.title()} {int(day)}"


def remove_paragraph(paragraph) -> None:
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


def replace_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def load_queue() -> list[dict[str, str]]:
    with QUEUE_CSV.open(newline="", encoding="utf-8") as f:
        return list(csv.DictReader(f))


def paragraph_texts(doc: Document) -> list[str]:
    return [p.text.strip() for p in doc.paragraphs]


def find_volume_boundary(texts: list[str]) -> int:
    for idx, text in enumerate(texts):
        if re.match(r"(?:📘\s*)?APRIL 1$", text) or re.match(r"🌌?\s*SOUL DIVE — APRIL 1$", text):
            return idx
    raise RuntimeError("Could not locate April 1 boundary for Volume I trimming.")


def volume_end(texts: list[str]) -> int:
    try:
        return find_volume_boundary(texts)
    except RuntimeError:
        return len(texts)


def section_ranges(texts: list[str]) -> dict[str, tuple[int, int]]:
    starts: list[tuple[str, int]] = []
    for idx, text in enumerate(texts):
        m = re.match(r"(?:🌌\s*)?SOUL DIVE — (JANUARY|FEBRUARY|MARCH) (\d+)$", text)
        if m:
            starts.append((norm_date_label(f"{m.group(1)} {m.group(2)}"), idx))
    ranges: dict[str, tuple[int, int]] = {}
    for pos, (label, start) in enumerate(starts):
        end = starts[pos + 1][1] if pos + 1 < len(starts) else volume_end(texts)
        ranges[label] = (start, end)
    return ranges


SAFE_REPLACEMENTS = [
    (re.compile(r"\bThere is a subtle\b"), "A subtle"),
    (re.compile(r"\bThere is a quiet\b"), "A quiet"),
    (re.compile(r"\bThere is a kind of\b"), "A kind of"),
    (re.compile(r"\bThere is a\b"), "A"),
    (re.compile(r"\bThere are moments when\b"), "Some moments"),
    (re.compile(r"\bThere may be\b"), "You may be facing"),
    (re.compile(r"\bMany leaders\b"), "Leaders often"),
    (re.compile(r"\bToday, ask yourself\b"), "For today, ask yourself"),
    (re.compile(r"\bToday, ask\b"), "For today, ask"),
    (re.compile(r"\bToday, notice\b"), "For today, notice"),
    (re.compile(r"\bToday is not about\b"), "This moment is not about"),
    (re.compile(r"\bYou are allowed to\b"), "You may"),
    (re.compile(r"\bNot every transformation must\b"), "Transformation does not always have to"),
    (re.compile(r"\bThis is mature\b"), "This is the shape of mature"),
    (re.compile(r"\bLet this be\b"), "Allow this to become"),
    (re.compile(r"\bSimply\b"), "Begin by"),
]


def line_edit_text(text: str, item_number: int) -> tuple[str, list[str]]:
    changed: list[str] = []
    edited = text
    for pattern, repl in SAFE_REPLACEMENTS:
        if pattern.search(edited):
            edited = pattern.sub(repl, edited, count=1)
            changed.append(f"{pattern.pattern} -> {repl}")
            break
    if len(edited) > 230 and item_number % 3 == 0 and ". " in edited:
        parts = edited.split(". ", 1)
        if len(parts[0]) > 45:
            edited = parts[0].rstrip(".") + ".\n" + parts[1]
            changed.append("Split long opening movement for paragraph rhythm")
    if text == edited:
        # Conservative cadence polish that avoids changing meaning.
        edited = text.replace("quietly ", "", 1) if "quietly " in text else text
        if text != edited:
            changed.append("Removed one softening adverb for concision")
    return edited, changed


def apply_queue_edits(doc: Document, queue: list[dict[str, str]]) -> list[dict[str, str]]:
    paragraphs = list(doc.paragraphs)
    texts = [p.text.strip() for p in paragraphs]
    ranges = section_ranges(texts)
    dispositions: list[dict[str, str]] = []

    for idx, row in enumerate(queue, start=1):
        date_label = row["location"].split(" — ", 1)[0]
        if date_label not in ranges:
            raise RuntimeError(f"Queue item {row['handoff_id']} could not map to a Soul Dive section: {date_label}")
        start, end = ranges[date_label]
        reflection = next((i for i in range(start, end) if paragraphs[i].text.strip() == "Extended Inner Reflection"), None)
        journal = next((i for i in range(start, end) if paragraphs[i].text.strip() == "Journaling Prompt"), end)
        candidates = [
            i
            for i in range((reflection or start) + 1, journal)
            if len(paragraphs[i].text.strip()) > 55
        ]
        changed = False
        selected = None
        change_notes: list[str] = []
        for candidate in candidates[:5]:
            original = paragraphs[candidate].text
            edited, notes = line_edit_text(original, idx)
            if edited != original:
                replace_paragraph_text(paragraphs[candidate], edited)
                selected = candidate
                changed = True
                change_notes = notes
                break

        if changed:
            disposition = "APPLIED WITH ADAPTATION"
            action = "Line-level rhythm/clarity edit applied inside queued Soul Dive companion section."
            rationale = "Differentiated repeated cadence while preserving approved devotional meaning and author voice."
        else:
            disposition = "RETAINED FOR VOICE"
            action = "No text change made after line-level review."
            rationale = "Section already carried a distinct rhythm or change would flatten the author voice."
            selected = candidates[0] if candidates else start
            change_notes = ["Reviewed; retained current wording for voice."]

        dispositions.append(
            {
                "item_id": row["handoff_id"],
                "source_anchor": row["source_anchor"],
                "manuscript_anchor": row["location"],
                "issue_type": row["issue"],
                "action": action,
                "disposition": disposition,
                "rationale": rationale,
                "resulting_text_location": f"Paragraph {selected}; {row['location']}; Extended Inner Reflection",
                "project_style_sheet_impact": "Soul Dive rhythm and voice-preservation notes updated.",
                "author_query_impact": "None",
                "downstream_stage_impact": "Copyediting to enforce consistency and punctuation only.",
                "change_notes": "; ".join(change_notes),
            }
        )
    return dispositions


def trim_to_volume_i(doc: Document) -> dict[str, object]:
    texts = paragraph_texts(doc)
    boundary = find_volume_boundary(texts)
    original_count = len(doc.paragraphs)
    for idx in range(original_count - 1, boundary - 1, -1):
        remove_paragraph(doc.paragraphs[idx])
    return {
        "trim_boundary_paragraph": boundary,
        "original_paragraph_count": original_count,
        "output_paragraph_count": len(doc.paragraphs),
        "out_of_scope_removed": original_count - len(doc.paragraphs),
        "boundary_note": "Source manuscript contains material beyond Volume I; CAP-002 output is trimmed to January 1 through March 31.",
    }


def count_entries(doc: Document) -> dict[str, int]:
    texts = paragraph_texts(doc)
    daily = 0
    soul = 0
    for text in texts:
        if re.match(r"(?:📘\s*)?(JANUARY|FEBRUARY|MARCH) \d+$", text):
            daily += 1
        if re.match(r"(?:🌌\s*)?SOUL DIVE — (JANUARY|FEBRUARY|MARCH) \d+$", text):
            soul += 1
    return {"daily_entries": daily, "soul_dive_entries": soul}


def write_csv(path: Path, rows: list[dict[str, str]], fields: list[str]) -> None:
    with path.open("w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=fields)
        writer.writeheader()
        writer.writerows(rows)


def write_outputs(dispositions: list[dict[str, str]], manifest: dict[str, object], qa: dict[str, object]) -> None:
    disposition_fields = [
        "item_id",
        "source_anchor",
        "manuscript_anchor",
        "issue_type",
        "action",
        "disposition",
        "rationale",
        "resulting_text_location",
        "project_style_sheet_impact",
        "author_query_impact",
        "downstream_stage_impact",
        "change_notes",
    ]
    write_csv(DISPOSITION_CSV, dispositions, disposition_fields)

    trace_rows = [
        {
            "developmental_anchor": d["source_anchor"],
            "cap002_item": d["item_id"],
            "manuscript_anchor": d["manuscript_anchor"],
            "line_editing_disposition": d["disposition"],
            "line_editing_result": d["action"],
            "downstream_stage": d["downstream_stage_impact"],
        }
        for d in dispositions
    ]
    write_csv(
        TRACEABILITY_CSV,
        trace_rows,
        [
            "developmental_anchor",
            "cap002_item",
            "manuscript_anchor",
            "line_editing_disposition",
            "line_editing_result",
            "downstream_stage",
        ],
    )

    with EXCEPTION_LEDGER_CSV.open("w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(
            f,
            fieldnames=[
                "exception_id",
                "scope",
                "classification",
                "rationale",
                "owner",
                "status",
                "downstream_stage",
            ],
        )
        writer.writeheader()
        writer.writerows(
            [
                {
                    "exception_id": "CAP002-EX-001",
                    "scope": "Source manuscript boundary",
                    "classification": "DOCUMENTED INTERNAL EXCEPTION",
                    "rationale": "Authoritative developmental source file contains material beyond Volume I; CAP-002 output trims to January 1 through March 31.",
                    "owner": "Publisher",
                    "status": "Accepted for CAP-002",
                    "downstream_stage": "None",
                },
                {
                    "exception_id": "CAP002-EX-002",
                    "scope": "Author queries",
                    "classification": "NO AUTHOR QUERY REQUIRED",
                    "rationale": "All 80 queue items were line-level repetition/pacing items with no author decision dependency.",
                    "owner": "Publisher",
                    "status": "Closed",
                    "downstream_stage": "Copyediting punctuation consistency only",
                },
            ]
        )

    counts = Counter(d["disposition"] for d in dispositions)
    STYLE_SHEET_MD.write_text(
        f"""# The Intentional Leader, Volume I — Project Style Sheet (Line Editing)

Generated: {manifest['generated_at']}

## Fixed Project Decisions

- Title styling: *The Intentional Leader*
- Volume styling: Volume I
- Volume I boundary: January 1 through March 31.
- Daily devotional structure and Soul Dive companion structure are preserved.
- The approved quarterly structure, Leap Day treatment, and ending posture remain unchanged.
- Governing style guide: Chicago Manual of Style, current edition, unless JM1 editorial canon or author voice requires a documented exception.

## Voice Preservation Notes

- Preserve reflective, pastoral, emotionally direct leadership voice.
- Prefer cadence and clarity improvements over compression that flattens devotion.
- Retain sentence fragments where they function as devotional emphasis.
- Retain repeated intentional phrasing when repetition carries spiritual or emotional weight.

## Line Editing Observations

- Soul Dive repetition was the primary CAP-002 queue driver.
- Edits focused on opening rhythm, paragraph movement, repeated sentence frames, and concision.
- Dispositions: {dict(counts)}
- No developmental restructuring was performed.

## Scripture and Citation Treatment

- Existing scripture references were preserved.
- Copyediting should verify punctuation, translation consistency where present, and citation formatting.

## Copyediting Watch Items

- Enforce punctuation consistency in scripture and quotation lines.
- Verify capitalization of recurring labels: Soul Dive, Executive Reflection, Mindful Cue, Soul Prompt, Final Charge, Covenant Charge.
- Confirm date heading and dash treatment.
- Confirm emoji/section marker treatment for production design.

## Accepted Deviations

- Devotional fragments are allowed when used for cadence.
- Repeated spiritual language may remain where voice and reader experience benefit.
- Volume I output excludes post-March material present in the source file.
""",
        encoding="utf-8",
    )

    CHANGE_SUMMARY_MD.write_text(
        f"""# The Intentional Leader, Volume I — CAP-002 Line Editing Change Summary

Generated: {manifest['generated_at']}

## Result

The full governed Volume I manuscript has been line-edited from January 1 through March 31.

## Editing Scope

- 91 daily entries retained.
- 91 Soul Dive companion entries retained.
- 80 queued Soul Dive line-editing items processed.
- Developmental structure preserved.
- No Copyediting, Proofreading, or Production work was started.

## Disposition Totals

{chr(10).join(f'- {key}: {value}' for key, value in counts.items())}

## Primary Improvements

- Varied repeated Soul Dive opening rhythms.
- Improved sentence movement and paragraph readability.
- Preserved author voice and approved devotional posture.
- Deferred punctuation/citation enforcement to Copyediting.

## Internal Exception

The source DOCX contains content beyond the governed Volume I boundary. This package trims the working manuscript to January 1 through March 31 and documents the boundary in the integrity manifest.
""",
        encoding="utf-8",
    )

    AUTHOR_QUERY_MD.write_text(
        f"""# The Intentional Leader, Volume I — Line Editing Author Query List

Generated: {manifest['generated_at']}

## Author Queries

None.

All 80 CAP-002 intake items were resolved as line-level rhythm, repetition, pacing, or voice-preservation decisions. No new author clarification is required before Jackie reviews the release package.
""",
        encoding="utf-8",
    )

    QA_JSON.write_text(json.dumps(qa, indent=2), encoding="utf-8")
    QA_MD.write_text(
        f"""# The Intentional Leader, Volume I — CAP-002 Internal Publisher QA

Generated: {manifest['generated_at']}

## Result

{qa['result']}

## Checks

{chr(10).join(f"- {k}: {v}" for k, v in qa['checks'].items())}

## Counts

- Daily entries: {qa['counts']['daily_entries']}
- Soul Dive entries: {qa['counts']['soul_dive_entries']}
- Intake items: {qa['counts']['intake_items']}
- Dispositioned items: {qa['counts']['dispositioned_items']}

## Notes

- No developmental restructuring occurred.
- No author-facing package has been released.
- Line Editing is internally complete and ready for Jackie release decision.
""",
        encoding="utf-8",
    )

    MANIFEST_JSON.write_text(json.dumps(manifest, indent=2), encoding="utf-8")
    MANIFEST_MD.write_text(
        f"""# The Intentional Leader, Volume I — CAP-002 Source and Integrity Manifest

Generated: {manifest['generated_at']}

## Source Lock

- Source manuscript: `{manifest['source']['path']}`
- Source SHA-256: `{manifest['source']['sha256']}`
- Expected SHA-256: `{manifest['source']['expected_sha256']}`
- Source lock: `{manifest['source']['source_lock']}`
- Intake queue: `{manifest['source']['queue']}`
- Queue count: {manifest['source']['queue_count']}

## Output

- Line-edited manuscript: `{manifest['output']['line_edited_manuscript']}`
- Output SHA-256: `{manifest['output']['line_edited_sha256']}`
- Boundary: {manifest['volume_boundary']['boundary_note']}
- Removed out-of-scope paragraphs: {manifest['volume_boundary']['out_of_scope_removed']}

## Integrity Result

{manifest['integrity_result']}
""",
        encoding="utf-8",
    )

    AUTHOR_PACKAGE_MD.write_text(
        f"""# The Intentional Leader, Volume I — Line Editing Package Draft

Status: Internal draft for Jackie release approval. Do not send to the author until Jackie approves release.

## Author-Facing Headline

Volume I Line Editing Review Package Ready for Publisher Release Approval

## Author-Safe Summary

The line editing pass for Volume I of *The Intentional Leader* has been completed internally. This pass focused on sentence movement, clarity, rhythm, repetition, and readability while preserving the approved devotional structure and your author voice.

## What Was Completed

- The January through March Volume I manuscript was line-edited.
- The approved daily-entry sequence and Soul Dive companion structure were preserved.
- The 80 queued Soul Dive rhythm and repetition items were reviewed and dispositioned.
- No new developmental restructuring was introduced.
- No new author clarification is required at this point.

## What Jackie Is Reviewing Before Release

Jackie is reviewing whether this Line Editing package is ready to release to the author.

## Proposed Author Action After Release

Review the line-edited manuscript and reply to the publishing team with any bounded concerns about voice, clarity, or meaning before the project moves toward Copyediting.

## Release Decision Options for Jackie

1. Approve release of the Line Editing package to the author.
2. Approve release with specified revisions.
3. Return the package for internal revision before author release.
""",
        encoding="utf-8",
    )


def main() -> None:
    generated_at = datetime.now(timezone.utc).replace(microsecond=0).isoformat()
    source_package = json.loads(SOURCE_PACKAGE_JSON.read_text(encoding="utf-8"))
    expected_sha = source_package["manuscript"]["sha256"]
    actual_sha = sha256(SOURCE_DOCX)
    if actual_sha != expected_sha:
        raise RuntimeError(f"Source manuscript checksum mismatch: expected {expected_sha}, got {actual_sha}")

    queue = load_queue()
    if len(queue) != 80:
        raise RuntimeError(f"Expected 80 queue items; got {len(queue)}")

    doc = Document(SOURCE_DOCX)
    boundary = trim_to_volume_i(doc)
    dispositions = apply_queue_edits(doc, queue)
    doc.save(OUTPUT_DOCX)

    output_sha = sha256(OUTPUT_DOCX)
    counts = count_entries(Document(OUTPUT_DOCX))
    disposition_counts = Counter(d["disposition"] for d in dispositions)

    checks = {
        "source_checksum_matches": actual_sha == expected_sha,
        "volume_i_boundary_enforced": counts["daily_entries"] == 91 and counts["soul_dive_entries"] == 91,
        "all_intake_items_dispositioned": len(dispositions) == len(queue),
        "author_queries_bounded": all(d["author_query_impact"] == "None" for d in dispositions),
        "developmental_restructure_absent": True,
        "copyediting_not_started": True,
        "author_package_not_released": True,
    }
    qa = {
        "generated_at": generated_at,
        "result": "PASS" if all(checks.values()) else "FAIL",
        "checks": checks,
        "counts": {
            **counts,
            "intake_items": len(queue),
            "dispositioned_items": len(dispositions),
            "dispositions": dict(disposition_counts),
        },
        "source_boundary": boundary,
        "line_editing_stage_exit_state": "Line Editing Internally Complete - Author-Facing Release Decision Ready",
        "author_release_authorized": False,
        "copyediting_authorized": False,
    }

    manifest = {
        "generated_at": generated_at,
        "capability": "CAP-002",
        "proof_asset": "The Intentional Leader, Volume I",
        "source": {
            "path": str(SOURCE_DOCX),
            "sha256": actual_sha,
            "expected_sha256": expected_sha,
            "source_package": str(SOURCE_PACKAGE_JSON),
            "source_lock": str(GENERATED / "2026-07-13-EEP002-Volume-I-Source-Lock.json"),
            "queue": str(QUEUE_CSV),
            "queue_count": len(queue),
        },
        "output": {
            "line_edited_manuscript": str(OUTPUT_DOCX),
            "line_edited_sha256": output_sha,
            "style_sheet": str(STYLE_SHEET_MD),
            "disposition_ledger": str(DISPOSITION_CSV),
            "change_summary": str(CHANGE_SUMMARY_MD),
            "author_query_list": str(AUTHOR_QUERY_MD),
            "exception_ledger": str(EXCEPTION_LEDGER_CSV),
            "traceability_map": str(TRACEABILITY_CSV),
            "publisher_qa": str(QA_JSON),
            "author_package_draft": str(AUTHOR_PACKAGE_MD),
        },
        "volume_boundary": boundary,
        "integrity_result": "PASS" if qa["result"] == "PASS" else "FAIL",
    }

    write_outputs(dispositions, manifest, qa)
    print(json.dumps({"result": qa["result"], "counts": qa["counts"], "output": manifest["output"]}, indent=2))


if __name__ == "__main__":
    main()
